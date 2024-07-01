import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from io import BytesIO
import base64
from streamlit_extras.switch_page_button import switch_page



def generate_docx(details, font_name):
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)

        # Add border to the document
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(ns.qn('w:val'), 'single')
            border.set(ns.qn('w:sz'), '4')
            border.set(ns.qn('w:space'), '24')
            border.set(ns.qn('w:color'), '000000')
            pgBorders.append(border)
        sectPr.append(pgBorders)

    # Add name
    name = doc.add_heading(level=1)
    run = name.add_run(details['name'])
    run.font.size = Pt(24)
    run.font.name = font_name
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add contact info
    contact_info = doc.add_paragraph()
    run = contact_info.add_run(f"Email: {details['email']} | Phone: {details['phone']}")
    run.font.name = font_name
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add sections
    doc.add_heading('Summary', level=2)
    summary_paragraph = doc.add_paragraph(details['summary'])
    for run in summary_paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

    if details['experience']:
        doc.add_heading('Experience', level=2)
        for exp in details['experience']:
            exp_paragraph = doc.add_paragraph()
            run = exp_paragraph.add_run(f"{exp['job_title']} at {exp['company']} ({exp['start_date']} - {exp['end_date']})\n")
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run = exp_paragraph.add_run(f"Responsibilities:\n{exp['responsibilities']}")
            run.font.name = font_name
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    if details['education']:
        doc.add_heading('Education', level=2)
        for edu in details['education']:
            edu_paragraph = doc.add_paragraph()
            run = edu_paragraph.add_run(f"{edu['degree']} - {edu['institution']} ({edu['graduation_date']})")
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_heading('Skills', level=2)
    skills_paragraph = doc.add_paragraph(details['skills'])
    for run in skills_paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_heading('Certifications', level=2)
    certifications_paragraph = doc.add_paragraph(details['certifications'])
    for run in certifications_paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

    return doc


def main():
    st.title("Resume Application")

    name = st.text_input("Name", placeholder="John Doe")
    email = st.text_input("Email", placeholder="john.doe@example.com")
    phone = st.text_input("Phone", placeholder="(123) 456-7890")
    summary = st.text_area("Summary", placeholder="A brief summary about yourself.")


    experience = []
    exp_count = st.number_input("Number of Experiences", min_value=1, max_value=7, step=1)
    for i in range(exp_count):
        with st.expander(f"Experience {i + 1}"):
            job_title = st.text_input(f"Job Title {i + 1}", placeholder="Software Engineer")
            company = st.text_input(f"Company {i + 1}", placeholder="ABC Corp")
            start_date = st.text_input(f"Start Date {i + 1}", placeholder="June 2023")
            end_date = st.text_input(f"End Date {i + 1}", placeholder="Present")
            responsibilities = st.text_area(f"Responsibilities {i + 1}", placeholder="List your job responsibilities.")
            experience.append({
                'job_title': job_title,
                'company': company,
                'start_date': start_date,
                'end_date': end_date,
                'responsibilities': responsibilities
            })

    education = []
    edu_count = st.number_input("Number of Education Entries", min_value=1, max_value=7, step=1)
    for i in range(edu_count):
        with st.expander(f"Education {i + 1}"):
            degree = st.text_input(f"Degree {i + 1}", placeholder="Bachelor of Science in Computer Science")
            institution = st.text_input(f"Institution {i + 1}", placeholder="XYZ University")
            graduation_date = st.text_input(f"Graduation Date {i + 1}", placeholder="June 2024")
            education.append({
                'degree': degree,
                'institution': institution,
                'graduation_date': graduation_date
            })

    skills = st.text_area("Skills", placeholder="List your skills.")
    certifications = st.text_area("Certifications", placeholder="List your certifications.")
    font_name = st.selectbox("Select Font", ["Arial", "Times New Roman", "Calibri", "Courier New", "Georgia"])

    if st.button("Generate Resume"):
        details = {
            'name': name,
            'email': email,
            'phone': phone,
            'summary': summary,
            'experience': experience,
            'education': education,
            'skills': skills,
            'certifications': certifications
        }
        doc = generate_docx(details, font_name)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        html = create_download_link(buffer.getvalue(), f"{name}.docx")
        st.markdown(html, unsafe_allow_html=True)

def create_download_link(val, filename):
    b64 = base64.b64encode(val).decode()  # val looks like b'...'
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">Download file</a>'

if __name__ == "__main__":
    main()
